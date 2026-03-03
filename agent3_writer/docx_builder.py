"""
Збірка .docx документа за стандартами українського судочинства
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_SECTION

from shared.config import settings
from shared.logger import get_logger
from shared.models import DocumentRequest

logger = get_logger(__name__)

DOCUMENT_TITLES = {
    "appeal": "АПЕЛЯЦІЙНА СКАРГА",
    "cassation": "КАСАЦІЙНА СКАРГА",
    "objection": "ВІДЗИВ",
    "motion_security": "КЛОПОТАННЯ про забезпечення позову",
    "motion_restore_deadline": "КЛОПОТАННЯ про поновлення строку",
    "motion_evidence": "КЛОПОТАННЯ про витребування доказів",
    "motion_expert": "КЛОПОТАННЯ про призначення експертизи",
    "motion_adjournment": "КЛОПОТАННЯ про відкладення розгляду",
}

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = Pt(18)  # ~1.5


class DocxBuilder:
    def build(self, document_text: str, request: DocumentRequest) -> str:
        """
        Створює .docx файл та повертає шлях до нього.
        """
        Path(settings.OUTPUT_PATH).mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_page(doc)

        parties = request.case_parties
        court_name = parties.get("court", "")
        doc_type_key = request.document_type.lower()

        self._add_header(doc, parties, court_name)
        self._add_title(doc, doc_type_key)
        self._add_case_number(doc, request.case_number)
        self._add_body(doc, document_text)
        self._add_appendices(doc, request.appendices)
        self._add_signature(doc, request.lawyer_name or "")
        self._add_disclaimer(doc)
        self._add_page_numbers(doc)

        # Ім'я файлу
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_num = request.case_number.replace("/", "-").replace(" ", "_")
        filename = f"{doc_type_key}_{safe_num}_{ts}.docx"
        output_path = Path(settings.OUTPUT_PATH) / filename

        doc.save(str(output_path))
        logger.info(f"Документ збережено: {output_path}")
        return str(output_path)

    # ------------------------------------------------------------------
    # Налаштування сторінки
    # ------------------------------------------------------------------

    def _setup_page(self, doc: Document) -> None:
        """Поля: ліво 3см, право 1.5см, верх 2см, низ 2см"""
        section = doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ------------------------------------------------------------------
    # Секції документа
    # ------------------------------------------------------------------

    def _add_header(self, doc: Document, parties: dict, court_name: str) -> None:
        """Шапка: до якого суду, від кого, до кого"""
        # Назва суду (по центру)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"До {court_name}")
        self._style_run(run, bold=False)

        # Сторони (вирівнювання по правому краю — типово для укр. практики)
        plaintiff = parties.get("plaintiff", "")
        defendant = parties.get("defendant", "")

        if plaintiff:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Апелянт/Позивач: {plaintiff}")
            self._style_run(run)

        if defendant:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Відповідач: {defendant}")
            self._style_run(run)

        doc.add_paragraph()  # Порожній рядок

    def _add_title(self, doc: Document, document_type: str) -> None:
        """Назва документа (по центру, жирний, 14pt)"""
        title_text = DOCUMENT_TITLES.get(document_type, document_type.upper())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        doc.add_paragraph()

    def _add_case_number(self, doc: Document, case_number: str) -> None:
        """Номер справи"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"у справі №{case_number}")
        self._style_run(run)
        doc.add_paragraph()

    def _add_body(self, doc: Document, text: str) -> None:
        """Основний текст (justify, відступ першого рядка 1.25см)"""
        for paragraph_text in text.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue

            # Заголовки секцій (рядки з великих літер або що починаються з цифри + крапки)
            lines = paragraph_text.splitlines()
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph()
                    continue

                p = doc.add_paragraph()
                is_heading = (
                    stripped.isupper()
                    or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)")
                )
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(stripped)
                self._style_run(run, bold=is_heading)

                if not is_heading:
                    # Відступ першого рядка
                    pPr = p._p.get_or_add_pPr()
                    ind = OxmlElement("w:ind")
                    ind.set(qn("w:firstLine"), "709")  # ~1.25см у одиницях EMU/20
                    pPr.append(ind)

                # Міжрядковий інтервал
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING

    def _add_appendices(self, doc: Document, appendices: list[str]) -> None:
        if not appendices:
            return
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Додатки:")
        self._style_run(run, bold=True)
        for i, app in enumerate(appendices, 1):
            p = doc.add_paragraph(f"{i}. {app}")
            self._style_run(p.runs[0])

    def _add_signature(self, doc: Document, lawyer_name: str) -> None:
        doc.add_paragraph()
        date_str = datetime.now().strftime("%d.%m.%Y")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{lawyer_name or 'Представник сторони'}\t\t{date_str}")
        self._style_run(run)

    def _add_disclaimer(self, doc: Document) -> None:
        """Footer disclaimer — як параграф наприкінці"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "* Документ згенеровано автоматично системою аналізу судової практики "
            "та потребує перевірки кваліфікованим юристом перед поданням до суду. *"
        )
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _add_page_numbers(self, doc: Document) -> None:
        """Нумерація сторінок знизу по центру"""
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    # ------------------------------------------------------------------
    # Утиліти
    # ------------------------------------------------------------------

    @staticmethod
    def _style_run(run, bold: bool = False) -> None:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.bold = bold
